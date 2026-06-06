import fitz
import pdfplumber
import docx
import openpyxl
import pandas as pd
import pytesseract
import cv2
import numpy as np


class DocumentService:
    async def _run_sync(self, func, *args, **kwargs):
        import asyncio
        return await asyncio.to_thread(func, *args, **kwargs)

    async def process(self, file_path: str, file_type: str) -> dict:
        processors = {
            "pdf": self._process_pdf,
            "docx": self._process_docx,
            "xlsx": self._process_xlsx,
            "csv": self._process_csv,
            "txt": self._process_text,
            "md": self._process_text,
            "json": self._process_text,
            "py": self._process_text,
            "js": self._process_text,
            "html": self._process_text,
            "css": self._process_text,
        }
        processor = processors.get(file_type, self._process_text)
        try:
            result = await processor(file_path)
            return result
        except Exception as e:
            return {"text": f"Error processing file: {str(e)}", "metadata": {}, "error": str(e)}

    async def _process_pdf(self, file_path: str) -> dict:
        import asyncio
        def _sync():
            text_parts = []
            metadata = {}
            with fitz.open(file_path) as doc:
                metadata = {
                    "pages": len(doc),
                    "title": doc.metadata.get("title", ""),
                    "author": doc.metadata.get("author", ""),
                }
                for page in doc:
                    text_parts.append(page.get_text())
            with pdfplumber.open(file_path) as pdf:
                tables = []
                for page in pdf.pages:
                    page_tables = page.extract_tables()
                    if page_tables:
                        tables.extend(page_tables)
                if tables:
                    text_parts.append("\n\n--- TABLES ---\n")
                    for t in tables:
                        df = pd.DataFrame(t)
                        text_parts.append(df.to_string())
            return {"text": "\n".join(text_parts), "metadata": metadata}
        return await asyncio.to_thread(_sync)

    async def _process_docx(self, file_path: str) -> dict:
        import asyncio
        def _sync():
            doc = docx.Document(file_path)
            text = "\n".join([p.text for p in doc.paragraphs])
            return {"text": text, "metadata": {"paragraphs": len(doc.paragraphs)}}
        return await asyncio.to_thread(_sync)

    async def _process_xlsx(self, file_path: str) -> dict:
        import asyncio
        def _sync():
            wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
            text_parts = []
            for sheet_name in wb.sheetnames:
                text_parts.append(f"\n=== Sheet: {sheet_name} ===")
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                text_parts.append(df.to_string())
            return {"text": "\n".join(text_parts), "metadata": {"sheets": wb.sheetnames}}
        return await asyncio.to_thread(_sync)

    async def _process_csv(self, file_path: str) -> dict:
        import asyncio
        def _sync():
            df = pd.read_csv(file_path)
            return {"text": df.to_string(), "metadata": {"rows": len(df), "columns": list(df.columns)}}
        return await asyncio.to_thread(_sync)

    async def _process_text(self, file_path: str) -> dict:
        import asyncio
        def _sync():
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                text = f.read()
            return {"text": text, "metadata": {"size": len(text)}}
        return await asyncio.to_thread(_sync)

    async def extract_images(self, file_path: str) -> list[str]:
        import asyncio
        def _sync():
            extracted_texts = []
            img = cv2.imread(file_path)
            if img is not None:
                gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
                _, thresh = cv2.threshold(gray, 150, 255, cv2.THRESH_BINARY)
                text = pytesseract.image_to_string(thresh)
                if text.strip():
                    extracted_texts.append(text)
            else:
                try:
                    from PIL import Image
                    img = Image.open(file_path)
                    text = pytesseract.image_to_string(img)
                    if text.strip():
                        extracted_texts.append(text)
                except Exception as e:
                    print(f"OCR fallback failed: {e}")
            return extracted_texts
        return await asyncio.to_thread(_sync)


class VisionService:
    async def analyze_image(self, image_path: str, prompt: str = "Describe this image in detail.") -> str:
        import base64
        with open(image_path, "rb") as f:
            image_data = base64.b64encode(f.read()).decode("utf-8")
        from app.core.llm import llm_service
        response = await llm_service.chat([
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": prompt},
                    {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{image_data}"}}
                ]
            }
        ])
        return response

    async def extract_text_from_image(self, image_path: str) -> str:
        try:
            img = cv2.imread(image_path)
            if img is None:
                from PIL import Image as PILImage
                img_pil = PILImage.open(image_path)
                img = cv2.cvtColor(np.array(img_pil), cv2.COLOR_RGB2BGR)

            from easyocr import Reader
            reader = Reader(["en"], gpu=False)
            results = reader.readtext(img)
            return "\n".join([text for _, text, conf in results if conf > 0.3])
        except Exception as e:
            return f"OCR failed: {str(e)}"


document_service = DocumentService()
vision_service = VisionService()
